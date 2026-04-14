from flask import Blueprint, jsonify, send_file
import os
import json
from .models import Document
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from io import BytesIO

bp = Blueprint('export_api', __name__)

@bp.route('/export/<document_id>/docx', methods=['GET'])
def export_docx(document_id: str):
    doc = Document.query.get(document_id)
    if not doc or doc.status != 'completed' or not doc.analysis:
        return jsonify({"error": "Document not found or analysis incomplete", "code": "NOT_READY"}), 404

    try:
        docx = DocxDocument()
        
        # Title
        title = docx.add_heading(f"Risk Analysis Report: {doc.filename}", 0)
        
        # Summary Section
        docx.add_heading("Executive Summary", level=1)
        docx.add_paragraph(doc.analysis.summary_text)

        # Clauses Section
        docx.add_heading("Key Identified Clauses", level=1)
        
        clauses = json.loads(doc.analysis.clauses_json)
        # Sort by risk level: Red > Yellow > Green
        risk_map = {"Red": 0, "Yellow": 1, "Green": 2}
        clauses.sort(key=lambda c: risk_map.get(c.get("category", "Yellow"), 3))

        for clause in clauses:
            cat = clause.get("category", "Yellow")
            if cat == "Green":
                continue # Optionally skip low risk, or include all. We'll include all but color them.
                
            heading = docx.add_heading(level=2)
            run = heading.add_run(f"[{cat}] {clause.get('type', 'Clause')} (Score: {clause.get('score', 0):.2f})")
            if cat == "Red":
                run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            elif cat == "Yellow":
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
            else:
                run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

            docx.add_paragraph(f"Text: {clause.get('text', '')}")
            docx.add_paragraph(f"Explanation: {clause.get('explanation', '')}").bold = True
            docx.add_paragraph("-" * 40)

        # Save to BytesIO stream
        f = BytesIO()
        docx.save(f)
        f.seek(0)
        
        return send_file(
            f,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{doc.filename}_Analysis.docx"
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": "Failed to generate DOCX report"}), 500
